import os
import google.generativeai as genai
import PyPDF2
from docx import Document
import requests
from bs4 import BeautifulSoup
from urllib.parse import urlparse
import tempfile

class DocumentSummarizer:
    def __init__(self):
        # Initialize Gemini for summarization
        genai.configure(api_key=os.getenv('GEMINI_API_KEY'))
        self.model = genai.GenerativeModel('gemini-2.5-flash')
    
    def extract_text_from_pdf(self, pdf_path):
        """Extract text from PDF file"""
        try:
            text = ""
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
            
            return text, num_pages
        except Exception as e:
            raise Exception(f"Error extracting PDF text: {str(e)}")
    
    def extract_text_from_docx(self, docx_path):
        """Extract text from DOCX file"""
        try:
            doc = Document(docx_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
            
            return text
        except Exception as e:
            raise Exception(f"Error extracting DOCX text: {str(e)}")
    
    def extract_text_from_txt(self, txt_path):
        """Extract text from TXT file"""
        try:
            with open(txt_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            raise Exception(f"Error reading TXT file: {str(e)}")
    
    def extract_text_from_url(self, url):
        """Extract text content from URL"""
        try:
            # Send request with headers to avoid blocking
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }
            response = requests.get(url, headers=headers, timeout=10)
            response.raise_for_status()
            
            # Parse HTML content
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract text
            text = soup.get_text()
            
            # Clean up text
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            # Get page title
            title = soup.find('title').string if soup.find('title') else "Untitled"
            
            return text, title
        except Exception as e:
            raise Exception(f"Error extracting URL content: {str(e)}")
    
    def generate_summary(self, text, content_type="document"):
        """Generate summary using Gemini"""
        try:
            # Truncate text if too long (Gemini has token limits)
            max_chars = 30000
            if len(text) > max_chars:
                text = text[:max_chars] + "...[content truncated]"
            
            prompt = f"""Please provide a comprehensive summary of the following {content_type}:

1. **Main Topics**: What are the key topics or themes covered?
2. **Key Points**: List the most important points or findings (bullet points)
3. **Executive Summary**: Provide a 2-3 paragraph executive summary
4. **Notable Insights**: Any particularly interesting or important insights
5. **Conclusions**: Main conclusions or takeaways

Content to summarize:
{text}

Please be thorough but concise, focusing on the most important information."""
            
            response = self.model.generate_content(prompt)
            
            return response.text if response.text else "Unable to generate summary."
            
        except Exception as e:
            return f"Summary generation failed: {str(e)}"
    
    def generate_brief_summary(self, text):
        """Generate a brief summary"""
        try:
            # Truncate text if too long
            max_chars = 10000
            if len(text) > max_chars:
                text = text[:max_chars] + "...[content truncated]"
            
            prompt = f"""Provide a brief 3-4 sentence summary of the following content, 
            highlighting only the most essential information:
            
            {text}"""
            
            response = self.model.generate_content(prompt)
            
            return response.text if response.text else "Unable to generate brief summary."
            
        except Exception as e:
            return f"Brief summary generation failed: {str(e)}"
    
    def extract_key_entities(self, text):
        """Extract key entities from text"""
        try:
            # Truncate text if too long
            max_chars = 10000
            if len(text) > max_chars:
                text = text[:max_chars]
            
            prompt = f"""Extract and list the following from the text:
            1. Key people/names mentioned
            2. Organizations/companies
            3. Locations/places
            4. Important dates/times
            5. Key concepts or technical terms
            
            Format as categorized lists.
            
            Text: {text}"""
            
            response = self.model.generate_content(prompt)
            
            return response.text if response.text else "No entities extracted."
            
        except Exception as e:
            return f"Entity extraction failed: {str(e)}"
    
    def summarize_document(self, file_path):
        """Main function to summarize documents"""
        try:
            # Determine file type and extract text
            file_ext = file_path.split('.')[-1].lower()
            
            if file_ext == 'pdf':
                text, num_pages = self.extract_text_from_pdf(file_path)
                metadata = {'type': 'PDF', 'pages': num_pages}
            elif file_ext in ['doc', 'docx']:
                text = self.extract_text_from_docx(file_path)
                metadata = {'type': 'Word Document'}
            elif file_ext == 'txt':
                text = self.extract_text_from_txt(file_path)
                metadata = {'type': 'Text File'}
            else:
                return {'error': f'Unsupported file type: {file_ext}'}
            
            if not text or len(text.strip()) < 10:
                return {'error': 'Document appears to be empty or contains no extractable text'}
            
            # Generate summaries
            detailed_summary = self.generate_summary(text, "document")
            brief_summary = self.generate_brief_summary(text)
            key_entities = self.extract_key_entities(text)
            
            return {
                'metadata': metadata,
                'word_count': len(text.split()),
                'character_count': len(text),
                'brief_summary': brief_summary,
                'detailed_summary': detailed_summary,
                'key_entities': key_entities
            }
            
        except Exception as e:
            return {'error': f'Document summarization failed: {str(e)}'}
    
    def summarize_url(self, url):
        """Summarize content from URL"""
        try:
            # Validate URL
            parsed = urlparse(url)
            if not parsed.scheme or not parsed.netloc:
                return {'error': 'Invalid URL format'}
            
            # Extract text from URL
            text, title = self.extract_text_from_url(url)
            
            if not text or len(text.strip()) < 10:
                return {'error': 'Unable to extract meaningful content from URL'}
            
            # Generate summaries
            detailed_summary = self.generate_summary(text, "webpage")
            brief_summary = self.generate_brief_summary(text)
            key_entities = self.extract_key_entities(text)
            
            return {
                'url': url,
                'title': title,
                'word_count': len(text.split()),
                'character_count': len(text),
                'brief_summary': brief_summary,
                'detailed_summary': detailed_summary,
                'key_entities': key_entities
            }
            
        except Exception as e:
            return {'error': f'URL summarization failed: {str(e)}'}
